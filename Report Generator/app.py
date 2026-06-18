#!/usr/bin/env python3
"""
Flask app for Completion Report Generator
"""

from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt
import os
from pathlib import Path
import json
from datetime import datetime

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max

# Services data
MAJOR_PROJECTS = {
    "full_reline": {
        "name": "Full Tank Reline",
        "brief": "Key Environmental Services were engaged to refurbish a cold-water storage tank which was suffering from corrosion and deterioration. We proposed that the internal surfaces be given long-term protection by the application of a solvent-free polyurethane coating system.",
        "before": "The tank showed evidence of heavy corrosion, rust, and delamination of the internal surfaces. Water quality was compromised and the structural integrity was at risk.",
        "prep": "The tank was drained and thoroughly dried. All corroded and affected internal surfaces were manually prepared to provide a surface profile suitable for coating application. Any holes or damage from corrosion were repaired using metal putty.",
        "basecoat": "All internal surfaces were coated with MAXLINE 100 DWPU at a nominal thickness of 500 microns, providing excellent adhesion and chemical resistance.",
        "topcoat": "The same coating system was applied for the top and final protective layer. The tank was then spray disinfected before being refilled and returned to service."
    },
    "bolt_reline": {
        "name": "Bolt restoration and lining",
        "brief": "Key Environmental Services were asked to refurbish corroded bolts and metal supports inside a cold-water storage tank. The bolts were suffering from corrosion in the oxygen-rich atmosphere above the waterline.",
        "before": "The bolts and metal supports were starting to corrode, with rust and staining present internally. Scale deposits were also evident around inlet areas.",
        "prep": "The tank was thoroughly dried and the affected bolts and support areas were manually prepared to provide the correct surface profile for coating.",
        "coating": "All corroded areas were brush coated with MAXLINE 100 DWPU at a nominal thickness of 500 microns. Once fully cured, the tank was spray disinfected and refilled."
    },
    "steel_reline": {
        "name": "Steel restorations and lining",
        "brief": "Key Environmental Services were commissioned to restore and protect steel components within a storage tank system. The steel was showing signs of rust and corrosion requiring protective coating.",
        "before": "Steel components showed significant rust formation and surface corrosion due to moisture and oxidation.",
        "prep": "Steel surfaces were cleaned and prepared to provide proper adhesion for the protective coating system.",
        "coating": "A MAXLINE 100 DWPU protective coating was applied to all steel components, providing long-term corrosion protection and chemical resistance."
    },
    "clean_disinfect": {
        "name": "Tank clean and disinfections",
        "brief": "Key Environmental Services were engaged to thoroughly clean and disinfect a storage tank to restore water quality and ensure microbiological safety.",
        "before": "The tank required deep cleaning due to sediment accumulation, biofilm growth, and water quality concerns.",
        "prep": "The tank was drained completely. Internal surfaces were scrubbed and cleaned to remove all debris, sediment, and biofilm deposits.",
        "coating": "Following cleaning, the tank was thoroughly disinfected using approved disinfection methods. The tank was then flushed, refilled with treated water, and sampled to confirm water quality compliance."
    },
    "tank_install": {
        "name": "Tank installation",
        "brief": "Key Environmental Services provided professional installation of a new water storage tank, including all necessary pipework connections and commissioning.",
        "before": "A new tank was required to replace or supplement existing storage capacity.",
        "prep": "The installation location was prepared with proper supports, level base, and access for maintenance.",
        "coating": "The new tank was installed, connected to supply and distribution pipework, tested for leaks, and commissioned into service with full system testing."
    },
    "tank_remove": {
        "name": "Tank removal",
        "brief": "Key Environmental Services safely removed and disposed of a redundant or damaged water storage tank in full compliance with environmental regulations.",
        "before": "An existing tank was no longer required or was beyond economic repair.",
        "prep": "The tank was isolated from the water system and fully drained.",
        "coating": "The tank was carefully removed, with all connections properly capped off. Appropriate disposal of the tank was arranged in line with waste regulations."
    },
    "crack_repair": {
        "name": "Repairing cracks/patch repairs on fibreglass tanks",
        "brief": "Key Environmental Services were engaged to repair structural cracks and damage to a fibreglass storage tank to restore its structural integrity and watertightness.",
        "before": "Visible cracks were present in the fibreglass structure, posing a risk of water leakage and structural failure.",
        "prep": "The affected areas were cleaned and prepared, with crack edges properly dressed to ensure good adhesion for the repair material.",
        "coating": "High-strength fibreglass repair resin was applied to all crack areas and patched sections. Once cured, repairs were sanded smooth and sealed with protective coating for durability."
    },
    "dividing_wall": {
        "name": "Dividing wall repairs on duplex fibreglass tanks",
        "brief": "Key Environmental Services repaired the internal dividing wall of a duplex fibreglass storage tank to restore separation between compartments and prevent cross-contamination.",
        "before": "The dividing wall showed cracks and deterioration, compromising the separation between tank compartments.",
        "prep": "Both sides of the dividing wall were cleaned and prepared to ensure proper adhesion of repair materials.",
        "coating": "Structural repairs were made using fibreglass reinforced resin on both sides of the wall. The wall was restored to full strength and sealed to prevent any water seepage between compartments."
    },
    "tanking_floor": {
        "name": "Tanking plant room floors",
        "brief": "Key Environmental Services applied waterproof tanking to the plant room floor to prevent water ingress and damage in case of tank or pipework leakage.",
        "before": "The floor lacked waterproof protection and was vulnerable to water damage from potential leaks.",
        "prep": "The floor surface was cleaned, dried, and any existing defects or cracks were repaired.",
        "coating": "A professional waterproof tanking system was applied to the entire floor area, creating a continuous protective layer with proper falls and drainage channels."
    }
}

REMEDIAL_WORKS = {
    "pipework_reroute": "Pipework rerouting",
    "vent_reroute": "Vent back rerouting",
    "pipework_disinfect": "Pipework disinfections",
    "outlet_flush": "Outlet flushing",
    "insulation_wrap": "Insulation wrapping pipework/tank",
    "lid_replace": "Lid replacement",
    "lid_vents": "Installation of lid vents and/or overflow screens",
    "support_replace": "Hollow support replacements",
    "fibreglass_repair": "Fibreglass repair and strengthening",
}

@app.route('/')
def index():
    # Convert project data for HTML (need to flatten it)
    projects_for_html = {k: v["name"] if isinstance(v, dict) else v
                        for k, v in MAJOR_PROJECTS.items()}
    return render_template('index.html',
                         major_projects=projects_for_html,
                         remedial_works=REMEDIAL_WORKS)

@app.route('/api/get-descriptions/<project_id>')
def get_descriptions(project_id):
    """Get pre-written descriptions for a project type"""
    if project_id in MAJOR_PROJECTS:
        project = MAJOR_PROJECTS[project_id]
        if isinstance(project, dict):
            return jsonify({
                'projectBrief': project.get('brief', ''),
                'before': project.get('before', ''),
                'preparation': project.get('prep', ''),
                'baseCoat': project.get('basecoat', ''),
                'topCoat': project.get('topcoat', ''),
                'coating': project.get('coating', ''),
            })
    return jsonify({'error': 'Project not found'}), 404

@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    """Generate completion report from form data"""
    try:
        data = request.form
        files = request.files

        # Extract form data
        job_details = {
            'project_number': data.get('projectNumber', ''),
            'customer_name': data.get('customerName', ''),
            'address_line1': data.get('addressLine1', ''),
            'address_line2': data.get('addressLine2', ''),
            'address_line3': data.get('addressLine3', ''),
            'postcode': data.get('postcode', ''),
            'completion_date': data.get('completionDate', ''),
        }

        descriptions = {
            'project_brief': data.get('projectBrief', ''),
            'before': data.get('before', ''),
            'preparation': data.get('preparation', ''),
            'base_coat': data.get('baseCoat', ''),
            'top_coat': data.get('topCoat', ''),
            'coating': data.get('coating', ''),
        }

        major_project = data.get('majorProject', '')
        work_completed_by = data.get('workCompletedBy', 'key_environmental')
        subcontractor_name = data.get('subcontractorName', '')
        remedial_works = data.getlist('remedialWorks[]')
        include_guarantee = data.get('includeGuarantee') == 'true'
        include_disinfection = data.get('includeDisinfection') == 'true'

        # Determine company name and footer info
        if work_completed_by == 'subcontracted' and subcontractor_name:
            company_name = subcontractor_name
            company_phone = "[Subcontractor Phone]"
            company_email = "[Subcontractor Email]"
            company_website = "[Subcontractor Website]"
            footer_line = f"For and on behalf of {subcontractor_name}"
        else:
            company_name = "Key Environmental Services Ltd"
            company_phone = "01789 330830"
            company_email = "service@key-environmental.co.uk"
            company_website = "www.watertankrelining.co.uk"
            footer_line = "For and on behalf of Key Environmental Services Ltd"

        # Convert remedial works IDs to names
        remedial_names = [REMEDIAL_WORKS.get(w, w) for w in remedial_works]

        # Load template
        template_path = 'completion_report_template.docx'
        if not os.path.exists(template_path):
            return jsonify({'error': 'Template file not found'}), 400

        doc = Document(template_path)

        # Replace placeholders
        replacements = {
            '[CUSTOMER_NAME]': job_details['customer_name'],
            '[ADDRESS_LINE_1]': job_details['address_line1'],
            '[ADDRESS_LINE_2]': job_details['address_line2'],
            '[ADDRESS_LINE_3]': job_details['address_line3'],
            '[POSTCODE]': job_details['postcode'],
            '[PROJECT_NUMBER]': job_details['project_number'],
            '[COMPLETION_DATE]': job_details['completion_date'],
            '[COMPANY_NAME]': company_name,
            '[COMPANY_PHONE]': company_phone,
            '[COMPANY_EMAIL]': company_email,
            '[COMPANY_WEBSITE]': company_website,
            '[PROJECT_BRIEF_DESCRIPTION]': descriptions['project_brief'],
            '[BEFORE_DESCRIPTION]': descriptions['before'],
            '[PREPARATION_DESCRIPTION]': descriptions['preparation'],
            '[BASE_COAT_DESCRIPTION]': descriptions.get('base_coat', ''),
            '[TOP_COAT_DESCRIPTION]': descriptions.get('top_coat', ''),
            '[COATING_DESCRIPTION]': descriptions.get('coating', ''),
            '[SIGNATURE_NAME]': 'CON O\'SHEA',
            '[SIGNATURE_TITLE]': 'DIRECTOR',
            '[FOOTER_LINE]': footer_line,
        }

        # Apply replacements to all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)

        # Add remedial works if applicable
        if remedial_names:
            remedial_text = "Remedial works completed: " + ", ".join(remedial_names)
            for paragraph in doc.paragraphs:
                if '[REMEDIAL_WORKS_DESCRIPTION]' in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace('[REMEDIAL_WORKS_DESCRIPTION]', remedial_text)

        # Handle photo uploads if present
        for section in ['before', 'preparation', 'baseCoat', 'topCoat', 'remedial']:
            file_key = f'photo_{section}'
            if file_key in files:
                photo_file = files[file_key]
                # Save photo temporarily and insert into doc
                if photo_file:
                    # This is a simplified version - you may want to embed photos properly
                    pass

        # Save document
        output_filename = f"{job_details['project_number']} - {job_details['customer_name']} - Completion Report.docx"
        output_path = os.path.join('reports', output_filename)

        # Create reports directory if it doesn't exist
        os.makedirs('reports', exist_ok=True)

        doc.save(output_path)

        return jsonify({
            'success': True,
            'filename': output_filename,
            'filepath': output_path
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<path:filepath>')
def download_report(filepath):
    """Download generated report"""
    try:
        return send_file(filepath, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
