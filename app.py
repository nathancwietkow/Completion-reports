#!/usr/bin/env python3
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

REMEDIAL_WORKS = {
    "pipework_reroute": "Pipework rerouting",
    "vent_reroute": "Vent back rerouting",
    "pipework_disinfect": "Pipework disinfections",
    "outlet_flush": "Outlet flushing",
    "insulation_wrap": "Insulation wrapping",
    "lid_replace": "Lid replacement",
    "lid_vents": "Lid vents/overflow screens",
    "support_replace": "Support replacements",
    "fibreglass_repair": "Fibreglass repair",
    "access_hatches": "Installing Access Hatches"
}

SUBCONTRACTORS = {
    "Echo Square Services Limited": "0800 7720 628"
}

SCOPE_OF_WORKS = {
    "Full Asset Reline": "We proposed that the tank(s) be given long term protection by the application of a solvent free polyurethane coating.\n\nThe product employed was [COATING], which offers numerous economic, technical and environmental features and benefits, some of which are highlighted below:\n\n• WRAS Approved\n• Complete Solvent Free Technology\n• Extremely Fast Curing Times, even at low temperatures\n• High degree of flexibility @ > 35%, capable of accommodating structural movement compared with resin / epoxy / bitumastic coatings of <1.0%\n• Superb adhesion to steel, concrete, GRP and many other substrates.\n• High levels of impact and chemical resistance.\n• Reduced downtime of storage tanks and process vessels - refill after minimum of 6hrs after completion.\n• Long life performance with minimal maintenance\n• Easy clean, ceramic tile-like finish\n• Application by brush / roller or plural component spray equipment\n• Excellent track record for applications such as drinking water tanks, cooling towers, reservoirs etc.",

    "Cooling Tower Reline": "We proposed that the tower(s) be given long term protection by the application of a solvent free polyurethane coating.\n\nThe product employed was [COATING], which offers numerous economic, technical and environmental features and benefits, some of which are highlighted below:\n\n• WRAS Approved\n• Complete Solvent Free Technology\n• Extremely Fast Curing Times, even at low temperatures\n• High degree of flexibility @ > 35%, capable of accommodating structural movement compared with resin / epoxy / bitumastic coatings of <1.0%\n• Superb adhesion to steel, concrete, GRP and many other substrates.\n• High levels of impact and chemical resistance.\n• Reduced downtime of storage tanks and process vessels - refill after minimum of 6hrs after completion.\n• Long life performance with minimal maintenance\n• Easy clean, ceramic tile-like finish\n• Application by brush / roller or plural component spray equipment\n• Excellent track record for applications such as drinking water tanks, cooling towers, reservoirs etc.",

    "Cooling Tower Repair": "We proposed that the tower(s) be given long term protection by the application of a solvent free polyurethane coating.\n\nThe product employed was [COATING], which offers numerous economic, technical and environmental features and benefits, some of which are highlighted below:\n\n• WRAS Approved\n• Complete Solvent Free Technology\n• Extremely Fast Curing Times, even at low temperatures\n• High degree of flexibility @ > 35%, capable of accommodating structural movement compared with resin / epoxy / bitumastic coatings of <1.0%\n• Superb adhesion to steel, concrete, GRP and many other substrates.\n• High levels of impact and chemical resistance.\n• Reduced downtime of storage tanks and process vessels - refill after minimum of 6hrs after completion.\n• Long life performance with minimal maintenance\n• Easy clean, ceramic tile-like finish\n• Application by brush / roller or plural component spray equipment\n• Excellent track record for applications such as drinking water tanks, cooling towers, reservoirs etc.",

    "Bolt restoration and lining": "We proposed that the tank(s) upper bolted areas be given long term protection by the application of a solvent free polyurethane coating.\n\nThe product employed was [COATING], which offers numerous economic, technical and environmental features and benefits, some of which are highlighted below:\n\n• WRAS Approved\n• Complete Solvent Free Technology\n• Extremely Fast Curing Times, even at low temperatures\n• High degree of flexibility @ > 35%, capable of accommodating structural movement compared with resin / epoxy / bitumastic coatings of <1.0%\n• Superb adhesion to steel, concrete, GRP and many other substrates.\n• High levels of impact and chemical resistance.\n• Reduced downtime of storage tanks and process vessels - refill after minimum of 6hrs after completion.\n• Long life performance with minimal maintenance\n• Easy clean, ceramic tile-like finish\n• Application by brush / roller or plural component spray equipment\n• Excellent track record for applications such as drinking water tanks, cooling towers, reservoirs etc.",

    "Steel restorations and lining": "We proposed that the tank(s) exposed steel areas be given long term protection by the application of a solvent free polyurethane coating.\n\nThe product employed was [COATING], which offers numerous economic, technical and environmental features and benefits, some of which are highlighted below:\n\n• WRAS Approved\n• Complete Solvent Free Technology\n• Extremely Fast Curing Times, even at low temperatures\n• High degree of flexibility @ > 35%, capable of accommodating structural movement compared with resin / epoxy / bitumastic coatings of <1.0%\n• Superb adhesion to steel, concrete, GRP and many other substrates.\n• High levels of impact and chemical resistance.\n• Reduced downtime of storage tanks and process vessels - refill after minimum of 6hrs after completion.\n• Long life performance with minimal maintenance\n• Easy clean, ceramic tile-like finish\n• Application by brush / roller or plural component spray equipment\n• Excellent track record for applications such as drinking water tanks, cooling towers, reservoirs etc.",

    "Asset clean and disinfections": "We proposed that the tanks would be drained, dried and thoroughly cleaned. These will then be disinfected, ahead of refilling, and reinstating back into the system.",

    "Repairing cracks/patch repairs": "We proposed that the areas of concern be reinforced, then given long term protection by the application of a solvent free polyurethane coating.\n\nThe product employed was [COATING], which offers numerous economic, technical and environmental features and benefits, some of which are highlighted below:\n\n• WRAS Approved\n• Complete Solvent Free Technology\n• Extremely Fast Curing Times, even at low temperatures\n• High degree of flexibility @ > 35%, capable of accommodating structural movement compared with resin / epoxy / bitumastic coatings of <1.0%\n• Superb adhesion to steel, concrete, GRP and many other substrates.\n• High levels of impact and chemical resistance.\n• Reduced downtime of storage tanks and process vessels - refill after minimum of 6hrs after completion.\n• Long life performance with minimal maintenance\n• Easy clean, ceramic tile-like finish\n• Application by brush / roller or plural component spray equipment\n• Excellent track record for applications such as drinking water tanks, cooling towers, reservoirs etc.",

    "Dividing wall repairs": "We proposed that the areas of concern be reinforced, then given long term protection by the application of a solvent free polyurethane coating.\n\nThe product employed was [COATING], which offers numerous economic, technical and environmental features and benefits, some of which are highlighted below:\n\n• WRAS Approved\n• Complete Solvent Free Technology\n• Extremely Fast Curing Times, even at low temperatures\n• High degree of flexibility @ > 35%, capable of accommodating structural movement compared with resin / epoxy / bitumastic coatings of <1.0%\n• Superb adhesion to steel, concrete, GRP and many other substrates.\n• High levels of impact and chemical resistance.\n• Reduced downtime of storage tanks and process vessels - refill after minimum of 6hrs after completion.\n• Long life performance with minimal maintenance\n• Easy clean, ceramic tile-like finish\n• Application by brush / roller or plural component spray equipment\n• Excellent track record for applications such as drinking water tanks, cooling towers, reservoirs etc.",

    "Tanking plant room floors": "We proposed that the areas of concern be reinforced, then given long term protection by the application of a solvent free polyurethane coating.\n\nThe product employed was [COATING], which offers numerous economic, technical and environmental features and benefits, some of which are highlighted below:\n\n• WRAS Approved\n• Complete Solvent Free Technology\n• Extremely Fast Curing Times, even at low temperatures\n• High degree of flexibility @ > 35%, capable of accommodating structural movement compared with resin / epoxy / bitumastic coatings of <1.0%\n• Superb adhesion to steel, concrete, GRP and many other substrates.\n• High levels of impact and chemical resistance.\n• Reduced downtime of storage tanks and process vessels - refill after minimum of 6hrs after completion.\n• Long life performance with minimal maintenance\n• Easy clean, ceramic tile-like finish\n• Application by brush / roller or plural component spray equipment\n• Excellent track record for applications such as drinking water tanks, cooling towers, reservoirs etc.",
}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    try:
        data = request.form

        # Collect ALL form values
        site_name = data.get('siteName', '')
        address_line_1 = data.get('addressLine1', '')
        address_line_2 = data.get('addressLine2', '')
        address_line_3 = data.get('addressLine3', '')
        postcode = data.get('postcode', '')
        asset_type = data.get('assetType', '')
        asset_subcategory = data.get('assetSubcategory', '')
        asset_location = data.get('assetLocation', '')
        asset_quantity = data.get('assetQuantity', '1')
        major_project = data.get('majorProject', '')
        project_number = data.get('projectNumber', '')
        completion_date = data.get('completionDate', '')
        company = data.get('workCompletedBy', 'Key Environmental Services Ltd')
        issues = data.getlist('assetIssues[]')

        # Set company contact details
        if company == "Key Environmental Services Ltd":
            company_contact_details = "01789 330830\nservice@key-environmental.co.uk\nwww.watertankrelining.co.uk"
        elif company in SUBCONTRACTORS:
            company_contact_details = SUBCONTRACTORS[company]
        else:
            company_contact_details = "[Contact details to be added]"

        # Get coating name
        coating_names = {
            'resichem_507': 'RESICHEM 507 DWPU',
            'maxline_100': 'MAXLINE 100 DWPU',
            'dwpu': 'DRINKING WATER POLYURETHANE (DWPU)'
        }
        coating_name = coating_names.get(data.get('coating', ''), 'the specified coating')

        # Get scope of works
        scope_of_works = SCOPE_OF_WORKS.get(major_project, '')
        scope_of_works = scope_of_works.replace('[COATING]', coating_name)

        if not os.path.exists('completion_report_template.docx'):
            return jsonify({'error': 'Template not found'}), 400

        doc = Document('completion_report_template.docx')

        # Get process step descriptions from form
        process_steps = {
            'before': data.get('processStepBefore', ''),
            'preparation': data.get('processStepPreparation', ''),
            'repairs': data.get('processStepRepairs', ''),
            'sealing': data.get('processStepSealing', ''),
            'basecoat': data.get('processStepBasecoat', ''),
            'topcoat': data.get('processStepTopcoat', ''),
            'refill': data.get('processStepRefill', ''),
            'other': data.get('processStepOther', ''),
        }

        # Build replacements dictionary
        replacements = {
            '[SITE_NAME]': site_name,
            '[ADDRESS_LINE_1]': address_line_1,
            '[ADDRESS_LINE_2]': address_line_2,
            '[ADDRESS_LINE_3]': address_line_3,
            '[POSTCODE]': postcode,
            '[ASSET_QUANTITY]': asset_quantity,
            '[ASSET_SUBCATEGORY]': asset_subcategory,
            '[ASSET_TYPE]': asset_type,
            '[ASSET_LOCATION]': asset_location,
            '[MAIN_PROJECT]': major_project,
            '[PROJECT_NUMBER]': project_number,
            '[COMPLETION_DATE]': completion_date,
            '[COMPANY_NAME]': company,
            '[COMPANY_CONTACT_DETAILS]': company_contact_details,
            '[SCOPE_OF_WORKS]': scope_of_works,
            '[BEFORE_DESCRIPTION]': process_steps['before'],
            '[PREPARATION_DESCRIPTION]': process_steps['preparation'],
            '[REPAIRS_DESCRIPTION]': process_steps['repairs'],
            '[SEALING_DESCRIPTION]': process_steps['sealing'],
            '[BASE_COAT_DESCRIPTION]': process_steps['basecoat'],
            '[TOP_COAT_DESCRIPTION]': process_steps['topcoat'],
            '[REFILL_DESCRIPTION]': process_steps['refill'],
        }

        # STEP 1: Replace all text placeholders (handles split placeholders across runs)
        def replace_paragraph_text(paragraph, replacements):
            """Replace all placeholders in a paragraph, even if split across runs."""
            full_text = ''.join(run.text for run in paragraph.runs)

            # Check if any replacement needed
            if not any(ph in full_text for ph in replacements):
                return

            # Apply all replacements
            new_text = full_text
            for placeholder, value in replacements.items():
                new_text = new_text.replace(placeholder, value)

            if new_text == full_text:
                return

            # Clear all runs and put new text in first run
            for run in paragraph.runs:
                run.text = ''

            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            replace_paragraph_text(paragraph, replacements)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_paragraph_text(paragraph, replacements)

        # Replace in headers/footers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                replace_paragraph_text(paragraph, replacements)

            for paragraph in section.footer.paragraphs:
                replace_paragraph_text(paragraph, replacements)

        # STEP 2: Handle [ISSUES_RAISED] - insert issue text into existing paragraphs
        issue_index = 0
        for paragraph in doc.paragraphs:
            if '[ISSUES_RAISED]' in paragraph.text:
                # Replace placeholder with actual issue
                if issue_index < len(issues):
                    for run in paragraph.runs:
                        if '[ISSUES_RAISED]' in run.text:
                            run.text = run.text.replace('[ISSUES_RAISED]', issues[issue_index])
                    issue_index += 1
                else:
                    # Remove extra [ISSUES_RAISED] if not enough issues
                    for run in paragraph.runs:
                        if '[ISSUES_RAISED]' in run.text:
                            run.text = run.text.replace('[ISSUES_RAISED]', '')

        # STEP 3: Handle photos - 2 column layout
        def insert_photo(doc, placeholder, photo_files, center=False):
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    # For front building photo, center align and add to existing paragraph
                    if center and placeholder == '[FRONT_BUILDING_PHOTO]':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Clear and add photo
                        for run in paragraph.runs:
                            run.text = ''

                        if photo_files and photo_files[0].filename:
                            photo_file = photo_files[0]
                            temp_path = f"/tmp/{photo_file.filename}"
                            photo_file.save(temp_path)
                            run = paragraph.add_run()
                            try:
                                run.add_picture(temp_path, width=Inches(3))
                            except:
                                run.text = "[Photo error]"
                            if os.path.exists(temp_path):
                                os.remove(temp_path)
                        return

                    # For other photos, use 2-column table layout
                    p_element = paragraph._element
                    parent = p_element.getparent()
                    insert_index = list(parent).index(p_element)

                    # Remove placeholder paragraph
                    parent.remove(p_element)

                    if photo_files and any(f.filename for f in photo_files):
                        num_photos = len([f for f in photo_files if f and f.filename])
                        num_rows = (num_photos + 1) // 2

                        table = doc.add_table(rows=num_rows, cols=2)
                        table.autofit = False

                        photo_index = 0
                        for row in table.rows:
                            for cell in row.cells:
                                if photo_index < len(photo_files):
                                    photo_file = photo_files[photo_index]
                                    if photo_file and photo_file.filename:
                                        temp_path = f"/tmp/{photo_file.filename}"
                                        photo_file.save(temp_path)
                                        cell.paragraphs[0].clear()
                                        run = cell.paragraphs[0].add_run()
                                        try:
                                            run.add_picture(temp_path, width=Inches(3))
                                        except:
                                            cell.paragraphs[0].text = "[Photo error]"
                                        if os.path.exists(temp_path):
                                            os.remove(temp_path)
                                    photo_index += 1

                        tbl = table._element
                        parent.insert(insert_index, tbl)
                    return

        # Insert photos
        photo_mapping = {
            'photo_front_building': ('[FRONT_BUILDING_PHOTO]', True),
            'photo_before': ('[BEFORE_PHOTO_PLACEHOLDER]', False),
            'photo_surface_prep': ('[PREPARATION_PHOTO_PLACEHOLDER]', False),
            'photo_basecoat': ('[BASE_COAT_PHOTO_PLACEHOLDER]', False),
            'photo_topcoat': ('[TOP_COAT_PHOTO_PLACEHOLDER]', False),
            'photo_remedials': ('[REMEDIAL_WORKS_PHOTO_PLACEHOLDER]', False),
        }

        for form_field, (placeholder, center) in photo_mapping.items():
            photos = request.files.getlist(form_field)
            if photos:
                insert_photo(doc, placeholder, photos, center=center)

        # STEP 4: Handle remedial works
        remedial = [REMEDIAL_WORKS.get(w, w) for w in data.getlist('remedialWorks[]')]
        if remedial:
            remedial_text = "Remedial works completed: " + ", ".join(remedial)
            for paragraph in doc.paragraphs:
                if '[REMEDIAL_WORKS_DESCRIPTION]' in paragraph.text:
                    for run in paragraph.runs:
                        if '[REMEDIAL_WORKS_DESCRIPTION]' in run.text:
                            run.text = run.text.replace('[REMEDIAL_WORKS_DESCRIPTION]', remedial_text)

        # Save
        filename = f"{project_number} - {site_name} - Completion Report.docx"
        path = os.path.join('reports', filename)
        os.makedirs('reports', exist_ok=True)
        doc.save(path)

        return jsonify({'success': True, 'filename': filename, 'filepath': path})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<path:filepath>')
def download(filepath):
    try:
        return send_file(filepath, as_attachment=True)
    except:
        return jsonify({'error': 'Download failed'}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
